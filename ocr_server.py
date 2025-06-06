from flask import Flask, request, jsonify, Response
from werkzeug.utils import secure_filename
import os
import pytesseract
import cv2
import numpy as np
from PIL import Image
import fitz  # PyMuPDF
import langid
import re
from spellchecker import SpellChecker
import json

app = Flask(__name__)

# Configuration
UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'png', 'jpg', 'jpeg', 'docx','doc', 'txt'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Set Tesseract path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Initialize spell checker
spell = SpellChecker()

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_text(text):
    if not text:
        return ""

    text = text.replace('\r\n', '\n').replace('\\n', '\n')
    text = re.sub(r'[ـ\xad]', '', text)
    text = re.sub(r'\u202b', '', text)
    text = re.sub(r'[ \t]+\n', '\n', text)
    text = re.sub(r'\n[ \t]+', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def spell_correct_text(text):
    corrected_words = []
    for word in text.split():
        if word.isalpha() and word.lower() not in spell:
            corrected = spell.correction(word)
            corrected_words.append(corrected if corrected else word)
        else:
            corrected_words.append(word)
    return ' '.join(corrected_words)

def preprocess_image_english(image):
    scale_percent = 150
    width = int(image.shape[1] * scale_percent / 100)
    height = int(image.shape[0] * scale_percent / 100)
    resized = cv2.resize(image, (width, height), interpolation=cv2.INTER_LINEAR)

    gray = cv2.cvtColor(resized, cv2.COLOR_BGR2GRAY)
    adaptive = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 11, 2
    )
    kernel = np.array([[0, -1, 0], [-1, 5, -1], [0, -1, 0]])
    sharpened = cv2.filter2D(adaptive, -1, kernel)
    denoised = cv2.fastNlMeansDenoising(sharpened, h=10)

    return denoised

def preprocess_image_arabic(image):
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    blurred = cv2.GaussianBlur(gray, (5, 5), 0)
    thresh = cv2.adaptiveThreshold(blurred, 255,
                                   cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                                   cv2.THRESH_BINARY, 11, 2)
    kernel = np.ones((2, 2), np.uint8)
    dilated = cv2.dilate(thresh, kernel, iterations=1)
    denoised = cv2.fastNlMeansDenoising(dilated, h=30)
    return denoised

def process_image(image_path):
    img = cv2.imread(image_path)
    if img is None:
        raise ValueError("Could not read image file")

    rough_text = pytesseract.image_to_string(img, config='--oem 3 --psm 6 -l eng+ara')
    lang, _ = langid.classify(rough_text)
    detected_lang = 'ara' if lang == 'ar' else 'eng'

    if detected_lang == 'ara':
        processed = preprocess_image_arabic(img)
        config = '--oem 3 --psm 6 -l ara'
    else:
        processed = preprocess_image_english(img)
        config = '--oem 3 --psm 6 -l eng -c user_defined_dpi=300'

    ocr_result = pytesseract.image_to_string(processed, config=config)

    if detected_lang == 'eng':
        ocr_result = spell_correct_text(ocr_result)

    return ocr_result

def process_pdf(pdf_path):
    pdf_document = fitz.open(pdf_path)
    full_text = ""

    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        pix = page.get_pixmap(matrix=fitz.Matrix(3.0, 3.0))
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)

        rough_text = pytesseract.image_to_string(img, config="--oem 3 --psm 6 -l eng+ara")
        lang, conf = langid.classify(rough_text)
        detected_lang = 'ara' if lang == 'ar' else 'eng'

        if detected_lang == 'ara':
            processed_img = preprocess_image_arabic(img)
            config = '--oem 3 --psm 11 -l ara'
        else:
            processed_img = preprocess_image_english(img)
            config = '--oem 3 --psm 6 -l eng -c user_defined_dpi=300'

        ocr_result = pytesseract.image_to_string(processed_img, config=config)

        if detected_lang == 'eng':
            ocr_result = spell_correct_text(ocr_result)

        full_text += f"--- Page {page_num + 1} ---\n{ocr_result}\n"

    pdf_document.close()
    return full_text

def process_docx(docx_path):
    from docx import Document
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        app.logger.error(f"DOCX processing error: {e}")
        return None

def process_doc(doc_path):
    try:
        import subprocess
        result = subprocess.run(['antiword', doc_path], capture_output=True, text=True)
        if result.returncode == 0:
            return result.stdout

        from doc import Document
        doc = Document(doc_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        app.logger.error(f"DOC processing error: {e}")
        return None

def process_text_file(txt_path):
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(txt_path, 'r', encoding='latin-1') as f:
            return f.read()

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)

        try:
            if file.filename.lower().endswith('.pdf'):
                result = process_pdf(file_path)
            elif file.filename.lower().endswith('.docx'):
                result = process_docx(file_path)
            elif file.filename.lower().endswith('.doc'):
                result = process_doc(file_path)
            elif file.filename.lower().endswith('.txt'):
                result = process_text_file(file_path)
            else:
                result = process_image(file_path)

            cleaned_text = clean_text(result)
            os.remove(file_path)

            if cleaned_text:
                return Response(
                    cleaned_text,
                    mimetype='text/plain; charset=utf-8',
                    headers={'Content-Language': 'en'}
                )
            return jsonify({"error": "Text extraction failed"}), 500

        except Exception as e:
            app.logger.error(f"Processing error: {e}")
            return jsonify({"error": str(e)}), 500

    return jsonify({"error": "File type not allowed"}), 400

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)